from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import importlib.util


class Resume_template:

    def __init__(
        self,
        name,
        address,
        city,
        state,
        number,
        email,
        objective,
        qualifications,
        profesional_experience,
        education,
        filename,
        document,
    ):

        self.name = name
        self.address = address
        self.city = city
        self.state = state
        self.number = number
        self.email = email
        self.objective = objective
        self.qualifications = qualifications
        self.profesional_experience = profesional_experience
        self.education = education
        self.filename = filename
        self.document = document

    def add_info(self):
        self.add_title()
        self.add_objective()
        self.add_qualifications()
        self.add_working_exp()
        self.add_education()
        self.document.save(self.filename)

    def add_title(self):
        title = self.document.add_paragraph()
        title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        first_and_last_name = title.add_run(self.name.title())
        first_and_last_name.font.size = Pt(22)

        personal_info = self.document.add_paragraph(
            f"{self.address.title()},\n Cell: {self.number}, Email: {self.email}"
        )
        personal_info.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_objective(self):
        obj_p = self.document.add_paragraph()
        run_1 = obj_p.add_run("Objective: ")
        run_1.bold = True
        run_1.font.size = Pt(14)
        obj_p.add_run(self.objective.capitalize())

    def add_working_exp(self):
        title = self.document.add_paragraph()
        run = title.add_run("Work Experience: ")
        run.bold = True
        run.underline = True
        run.font.size = Pt(14)
        for job in self.profesional_experience:
            if job["job"]["company_name"] != "":

                place_worked = job["job"]["company_name"]
                city = job["job"]["city"]
                state = job["job"]["state"]
                starting_date = job["job"]["starting_date"]
                end_date = job["job"]["end_date"]
                job_title = job["job"]["job_title"]
                job_desc = [
                    job["job"]["job_description"][0],
                    job["job"]["job_description"][1],
                    job["job"]["job_description"][2],
                ]

                title = self.document.add_paragraph()

                run_1 = title.add_run(f"{job_title.title()}")
                run_1.bold = True

                job_location = self.document.add_paragraph()

                run_1 = job_location.add_run(
                    f"{place_worked.capitalize()} - {city.capitalize()} {state.capitalize()}"
                )

                work_dates = self.document.add_paragraph()

                run_1 = work_dates.add_run(f"{starting_date} - {end_date}")

                job_des = self.document.add_paragraph()
                job_des1 = job_des.add_run(f"{job_desc[0].capitalize()}")
                job_des = self.document.add_paragraph()
                job_des2 = job_des.add_run(f"{job_desc[1].capitalize()}")
                job_des = self.document.add_paragraph()
                job_des2 = job_des.add_run(f"{job_desc[2].capitalize()}")

    def add_education(self):
        title = self.document.add_paragraph()
        run_1 = title.add_run("Education:")
        run_1.bold = True
        run_1.underline = True
        run_1.font.size = Pt(14)

        self.document.add_paragraph(self.education.capitalize())

    def add_qualifications(self):
        title = self.document.add_paragraph()
        run_1 = title.add_run("Qualifications:")
        run_1.bold = True
        run_1.underline = True
        run_1.font.size = Pt(14)

        for skill in self.qualifications:
            self.document.add_paragraph(skill.capitalize(), style="List Bullet")
